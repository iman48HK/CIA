import httpx
import json
import base64
import binascii
import math
import re
from datetime import datetime, timezone
from typing import Any, Literal
from html import escape
from pathlib import Path
from xml.sax.saxutils import escape as xml_escape
from uuid import uuid4

from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException
from fastapi.responses import FileResponse

from app.config import settings
from app.database import SessionLocal, get_db
from app.deps import get_current_user
from app.models import (
    OrdinanceFile,
    Project,
    ProjectDrawingUpload,
    ProjectFileUpload,
    ProjectOrdinanceSelection,
    ReportJob,
    User,
)
from app.schemas import AIChatRequest, AIChatResponse
from pydantic import BaseModel, Field
from sqlalchemy.orm import Session
import markdown as md

router = APIRouter(prefix="/ai", tags=["ai"])

OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"
REPORTS_DIR = Path(__file__).resolve().parents[2] / "storage" / "reports"
REPORT_BRAND_HEADER = "Construction Insight Agent"
# One stack for HTML/PDF/Word: common on Windows + Linux PDF engines; avoids mixed serif/mono in exports.
REPORT_FONT_STACK = '"Segoe UI", Arial, Helvetica, "Liberation Sans", sans-serif'
REPORT_FONT_STACK_XML = "Segoe UI, Arial, Helvetica, Liberation Sans, sans-serif"

# Report PDF viewport width (px): matches .wrap max-width 920 + horizontal padding in report CSS.
_REPORT_PDF_PAGE_WIDTH_PX = 992
_REPORT_PDF_MAX_HEIGHT_PX = 50_000


def _html_to_pdf_bytes_xhtml2pdf(html: str) -> bytes:
    from io import BytesIO

    from xhtml2pdf import pisa

    buf = BytesIO()
    pisa.CreatePDF(BytesIO(html.encode("utf-8")), dest=buf, encoding="utf-8")
    return buf.getvalue()


def _trim_trailing_blank_pdf_pages(pdf_bytes: bytes) -> bytes:
    """Chromium sometimes emits an extra empty page for custom paper sizes; drop trailing pages with no text."""
    import fitz

    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception:
        return pdf_bytes
    try:
        while doc.page_count > 1:
            last = doc[-1]
            if last.get_text().strip() or last.get_images():
                break
            doc.delete_page(doc.page_count - 1)
        return doc.tobytes()
    finally:
        doc.close()


def _html_to_pdf_bytes_chromium(html: str) -> bytes:
    """Print the same HTML the user downloads, using headless Chromium (screen media, full-page height)."""
    from playwright.sync_api import sync_playwright

    w = _REPORT_PDF_PAGE_WIDTH_PX
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        try:
            page = browser.new_page(viewport={"width": w, "height": min(1600, _REPORT_PDF_MAX_HEIGHT_PX)})
            page.emulate_media(media="screen")
            page.set_content(html, wait_until="load", timeout=120_000)
            try:
                page.evaluate("() => document.fonts.ready")
            except Exception:
                pass
            page.wait_for_timeout(200)
            height_px = int(
                page.evaluate(
                    """() => Math.ceil(Math.max(
                        document.body ? document.body.scrollHeight : 0,
                        document.documentElement ? document.documentElement.scrollHeight : 0
                    ))"""
                )
            )
            height_px = max(min(height_px + 40, _REPORT_PDF_MAX_HEIGHT_PX), 400)
            raw = page.pdf(
                print_background=True,
                width=f"{w}px",
                height=f"{height_px}px",
                margin={"top": "0", "right": "0", "bottom": "0", "left": "0"},
            )
            return _trim_trailing_blank_pdf_pages(raw)
        finally:
            browser.close()


def _html_to_pdf_bytes_for_report(html: str) -> bytes:
    if settings.report_pdf_chromium:
        try:
            return _html_to_pdf_bytes_chromium(html)
        except Exception:
            pass
    return _html_to_pdf_bytes_xhtml2pdf(html)


def _apply_report_docx_base_font(document: object) -> None:
    """Align Word export with report HTML sans stack (Arial matches PDF/HTML fallbacks)."""
    from docx.shared import Pt

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)


class ProjectAIChatRequest(AIChatRequest):
    project_id: int


class ConsolidatedReportRequest(BaseModel):
    project_id: int
    user_prompts: list[str] = Field(default_factory=list, max_length=20)
    chat_history: list[dict] = Field(default_factory=list, max_length=500)
    annotation_assets: list[dict] = Field(default_factory=list, max_length=30)
    author: str | None = Field(default=None, max_length=200)


class AnnotationMarkerIn(BaseModel):
    xPct: float = Field(ge=0, le=1)
    yPct: float = Field(ge=0, le=1)
    type: str = Field(default="deficiency", max_length=20)
    note: str = Field(default="", max_length=200)


class AnnotationAssetIn(BaseModel):
    drawing_id: int
    source_kind: Literal["drawing", "project_file"] = "drawing"
    filename: str = Field(min_length=1, max_length=255)
    image_data_url: str = Field(min_length=1, max_length=6_000_000)
    markers: list[AnnotationMarkerIn] = Field(default_factory=list, max_length=500)


class CollectedMaterialIn(BaseModel):
    id: str = Field(default="", max_length=80)
    source: str = Field(default="assistant", max_length=120)
    label: str = Field(default="Collected output", max_length=255)
    text: str = Field(min_length=1, max_length=200_000)
    created_at: str = Field(default="", max_length=80)


class ConsolidateCollectionReportRequest(BaseModel):
    project_id: int
    items: list[CollectedMaterialIn] = Field(default_factory=list, min_length=1, max_length=80)
    chat_history: list[dict] = Field(default_factory=list, max_length=500)
    annotation_assets: list[AnnotationAssetIn] = Field(default_factory=list, max_length=30)
    include_annotated_drawing: bool = False
    author: str | None = Field(default=None, max_length=200)


class CustomReportJobAccepted(BaseModel):
    job_id: int
    status: str = "pending"


class CustomReportJobStatusOut(BaseModel):
    job_id: int
    status: str
    report_type: str | None = None
    report_id: str | None = None
    downloads: dict[str, str] | None = None
    error_message: str | None = None


class StandardReportJobAccepted(BaseModel):
    job_id: int
    status: str = "pending"


class StandardReportJobStatusOut(BaseModel):
    job_id: int
    status: str
    report_type: str | None = None
    report_id: str | None = None
    downloads: dict[str, str] | None = None
    error_message: str | None = None


class ChatJobAccepted(BaseModel):
    job_id: int
    status: str = "pending"


class ChatJobStatusOut(BaseModel):
    job_id: int
    status: str
    reply: str | None = None
    error_message: str | None = None


async def _request_openrouter(model: str, message: str, headers: dict[str, str]) -> httpx.Response:
    system = (
        "You are AI Assistant for construction projects. "
        "Return responses in clean Markdown with this structure:\n"
        "## Quick Summary\n"
        "- 2-4 concise bullets\n\n"
        "## Key Findings\n"
        "- Ordered by impact\n\n"
        "## Recommended Actions\n"
        "- Actionable next steps\n\n"
        "## References / Citations\n"
        "- Cite only project-scoped sources: Drawings, Project Files, and Selected Ordinance Docs\n\n"
        "## Confidence\n"
        "- confidence_score: <0-100>\n"
        "- doubt_score: <0-100>\n"
        "- manual_review_reason: <why review is needed or 'None'>\n\n"
        "Rules:\n"
        "- Be concise and practical.\n"
        "- If data is missing, explicitly state assumptions.\n"
        "- Use only project documents as references (Drawings / Project Files / Selected Ordinance Docs).\n"
        "- Never cite outside sources or general web/standards references.\n"
        "- Never invent citations; if unavailable, say so using project document gaps only.\n"
        "- Use clear headings and bullet points only (no long paragraphs)."
    )
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": message},
        ],
    }
    async with httpx.AsyncClient(timeout=120.0) as client:
        return await client.post(OPENROUTER_URL, json=payload, headers=headers)


def _request_openrouter_sync(model: str, message: str, headers: dict[str, str]) -> httpx.Response:
    system = (
        "You are AI Assistant for construction projects. "
        "Return responses in clean Markdown with this structure:\n"
        "## Quick Summary\n"
        "- 2-4 concise bullets\n\n"
        "## Key Findings\n"
        "- Ordered by impact\n\n"
        "## Recommended Actions\n"
        "- Actionable next steps\n\n"
        "## References / Citations\n"
        "- Cite only project-scoped sources: Drawings, Project Files, and Selected Ordinance Docs\n\n"
        "## Confidence\n"
        "- confidence_score: <0-100>\n"
        "- doubt_score: <0-100>\n"
        "- manual_review_reason: <why review is needed or 'None'>\n\n"
        "Rules:\n"
        "- Be concise and practical.\n"
        "- If data is missing, explicitly state assumptions.\n"
        "- Use only project documents as references (Drawings / Project Files / Selected Ordinance Docs).\n"
        "- Never cite outside sources or general web/standards references.\n"
        "- Never invent citations; if unavailable, say so using project document gaps only.\n"
        "- Use clear headings and bullet points only (no long paragraphs)."
    )
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": message},
        ],
    }
    with httpx.Client(timeout=120.0) as client:
        return client.post(OPENROUTER_URL, json=payload, headers=headers)


def _project_context_message(project: Project, db: Session) -> str:
    drawing_rows = (
        db.query(ProjectDrawingUpload)
        .filter(ProjectDrawingUpload.project_id == project.id)
        .order_by(ProjectDrawingUpload.id.desc())
        .all()
    )
    selected_ordinance_ids = [
        oid
        for (oid,) in db.query(ProjectOrdinanceSelection.ordinance_file_id)
        .filter(ProjectOrdinanceSelection.project_id == project.id)
        .all()
    ]
    ordinance_titles = [
        title
        for (title,) in db.query(OrdinanceFile.title)
        .filter(OrdinanceFile.id.in_(selected_ordinance_ids))
        .all()
    ] if selected_ordinance_ids else []
    project_file_names = [
        fname
        for (fname,) in db.query(ProjectFileUpload.filename)
        .filter(ProjectFileUpload.project_id == project.id)
        .order_by(ProjectFileUpload.id.desc())
        .all()
    ]
    return (
        f"[Project: {project.name} | Project ID: {project.id}]\n"
        f"[Drawings: {', '.join([d.filename for d in drawing_rows]) if drawing_rows else 'None'}]\n"
        f"[Selected Ordinance Docs: {', '.join(ordinance_titles) if ordinance_titles else 'None'}]\n"
        f"[Optional Project Files: {', '.join(project_file_names) if project_file_names else 'None'}]\n"
    )


def _run_selected_prompts(
    prompts: list[str],
    project: Project,
    db: Session,
) -> list[dict]:
    clean_prompts = [p.strip() for p in prompts if isinstance(p, str) and p.strip()]
    if not clean_prompts:
        return []
    if not settings.openrouter_api_key:
        return []
    headers = {
        "Authorization": f"Bearer {settings.openrouter_api_key}",
        "HTTP-Referer": "http://localhost:3000",
        "X-Title": "Construction Insight Agent",
        "Content-Type": "application/json",
    }
    prefix = _project_context_message(project, db)
    transcript: list[dict] = []
    for prompt in clean_prompts:
        scoped = f"{prefix}{prompt}"
        try:
            r = _request_openrouter_sync(settings.openrouter_model, scoped, headers)
            if r.status_code != 200 and settings.openrouter_fallback_model:
                r = _request_openrouter_sync(settings.openrouter_fallback_model, scoped, headers)
            if r.status_code != 200:
                reply = f"Unable to generate AI response for this prompt. ({r.status_code})"
            else:
                data = r.json()
                reply = str(data["choices"][0]["message"]["content"] or "")
        except Exception:
            reply = "Unable to generate AI response for this prompt."
        now = datetime.now(timezone.utc).isoformat()
        transcript.append({"role": "user", "text": prompt, "created_at": now})
        transcript.append({"role": "assistant", "text": reply, "created_at": now})
    return transcript


@router.post("/chat", response_model=AIChatResponse)
async def chat(
    body: ProjectAIChatRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    reply = _run_single_chat_reply(body.project_id, user.id, body.message, db)
    return AIChatResponse(reply=reply)


@router.post("/chat/jobs", response_model=ChatJobAccepted)
def enqueue_chat_job(
    body: ProjectAIChatRequest,
    background_tasks: BackgroundTasks,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    _validate_chat_prereqs(body.project_id, user.id, db)
    job = ReportJob(
        owner_id=user.id,
        project_id=body.project_id,
        kind="chat_reply",
        status="pending",
    )
    db.add(job)
    db.commit()
    db.refresh(job)
    background_tasks.add_task(_run_chat_job_task, job.id, body.project_id, user.id, body.message)
    return ChatJobAccepted(job_id=job.id, status="pending")


@router.get("/chat/jobs/{job_id}", response_model=ChatJobStatusOut)
def get_chat_job_status(
    job_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    job = (
        db.query(ReportJob)
        .filter(ReportJob.id == job_id, ReportJob.owner_id == user.id, ReportJob.kind == "chat_reply")
        .first()
    )
    if not job:
        raise HTTPException(status_code=404, detail="Chat job not found")
    payload = json.loads(job.downloads_json) if job.downloads_json else {}
    reply = payload.get("reply") if isinstance(payload, dict) else None
    return ChatJobStatusOut(
        job_id=job.id,
        status=job.status,
        reply=reply if isinstance(reply, str) else None,
        error_message=job.error_message,
    )


def _validate_chat_prereqs(project_id: int, owner_id: int, db: Session) -> Project:
    project = db.query(Project).filter(Project.id == project_id, Project.owner_id == owner_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    drawing_rows = (
        db.query(ProjectDrawingUpload)
        .filter(ProjectDrawingUpload.project_id == project.id)
        .order_by(ProjectDrawingUpload.id.desc())
        .all()
    )
    if not drawing_rows:
        raise HTTPException(
            status_code=400,
            detail="Selected project must have at least one drawing before asking CIA Assistant.",
        )
    selected_ordinance_ids = [
        oid
        for (oid,) in db.query(ProjectOrdinanceSelection.ordinance_file_id)
        .filter(ProjectOrdinanceSelection.project_id == project.id)
        .all()
    ]
    if not selected_ordinance_ids:
        raise HTTPException(
            status_code=400,
            detail="Select at least one ordinance document for this project before asking CIA Assistant.",
        )
    if not settings.openrouter_api_key:
        raise HTTPException(
            status_code=503,
            detail="OpenRouter API key not configured. Set OPENROUTER_API_KEY in backend .env",
        )
    return project


def _run_single_chat_reply(project_id: int, owner_id: int, message: str, db: Session) -> str:
    project = _validate_chat_prereqs(project_id, owner_id, db)
    drawing_rows = (
        db.query(ProjectDrawingUpload)
        .filter(ProjectDrawingUpload.project_id == project.id)
        .order_by(ProjectDrawingUpload.id.desc())
        .all()
    )
    selected_ordinance_ids = [
        oid
        for (oid,) in db.query(ProjectOrdinanceSelection.ordinance_file_id)
        .filter(ProjectOrdinanceSelection.project_id == project.id)
        .all()
    ]
    ordinance_titles = [
        title
        for (title,) in db.query(OrdinanceFile.title)
        .filter(OrdinanceFile.id.in_(selected_ordinance_ids))
        .all()
    ]
    project_file_names = [
        fname
        for (fname,) in db.query(ProjectFileUpload.filename)
        .filter(ProjectFileUpload.project_id == project.id)
        .order_by(ProjectFileUpload.id.desc())
        .all()
    ]
    headers = {
        "Authorization": f"Bearer {settings.openrouter_api_key}",
        "HTTP-Referer": "http://localhost:3000",
        "X-Title": "Construction Insight Agent",
        "Content-Type": "application/json",
    }
    scoped_message = (
        f"[Project: {project.name} | Project ID: {project.id}]\n"
        f"[Drawings: {', '.join([d.filename for d in drawing_rows])}]\n"
        f"[Selected Ordinance Docs: {', '.join(ordinance_titles)}]\n"
        f"[Optional Project Files: {', '.join(project_file_names) if project_file_names else 'None'}]\n"
        f"{message}"
    )
    try:
        r = _request_openrouter_sync(settings.openrouter_model, scoped_message, headers)
        if r.status_code != 200 and settings.openrouter_fallback_model:
            r = _request_openrouter_sync(settings.openrouter_fallback_model, scoped_message, headers)
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"OpenRouter request failed: {e}") from e
    if r.status_code != 200:
        raise HTTPException(status_code=502, detail=f"OpenRouter error: {r.status_code} {r.text[:500]}")
    data = r.json()
    try:
        reply = str(data["choices"][0]["message"]["content"] or "")
    except (KeyError, IndexError) as e:
        raise HTTPException(status_code=502, detail="Unexpected OpenRouter response") from e
    return reply


def _run_chat_job_task(job_id: int, project_id: int, owner_id: int, message: str) -> None:
    db = SessionLocal()
    try:
        job = (
            db.query(ReportJob)
            .filter(ReportJob.id == job_id, ReportJob.owner_id == owner_id, ReportJob.kind == "chat_reply")
            .first()
        )
        if not job:
            return
        job.status = "running"
        db.commit()
        reply = _run_single_chat_reply(project_id, owner_id, message, db)
        job.status = "completed"
        job.downloads_json = json.dumps({"reply": reply})
        job.completed_at = datetime.now(timezone.utc)
        db.commit()
    except Exception as e:
        db.rollback()
        job = db.query(ReportJob).filter(ReportJob.id == job_id).first()
        if job:
            job.status = "failed"
            job.error_message = (str(e) or "Chat job failed")[:4000]
            job.completed_at = datetime.now(timezone.utc)
            db.commit()
    finally:
        db.close()


def _normalize_chat_text(s: str) -> str:
    return " ".join(s.strip().split())


def _assistant_reply_after_user_message(chat_history: list[dict], user_text: str) -> str:
    """Return the assistant message immediately following the first matching user turn."""
    target = user_text.strip()
    target_norm = _normalize_chat_text(user_text)
    n = len(chat_history)
    for i, row in enumerate(chat_history):
        if str(row.get("role", "")).lower() != "user":
            continue
        ut = str(row.get("text", "")).strip()
        if ut != target and _normalize_chat_text(ut) != target_norm:
            continue
        for j in range(i + 1, n):
            role = str(chat_history[j].get("role", "")).lower()
            if role == "assistant":
                return str(chat_history[j].get("text", "")).strip()
            if role == "user":
                break
    return "No assistant reply appears after this message in the exported transcript."


_HEADING_RE = re.compile(r"^\s{0,3}#{1,6}\s+(.+?)\s*$")


def _parse_markdown_sections(markdown_text: str) -> dict[str, str]:
    sections: dict[str, list[str]] = {}
    current = "General"
    sections[current] = []
    for raw_line in str(markdown_text or "").splitlines():
        m = _HEADING_RE.match(raw_line)
        if m:
            current = m.group(1).strip()
            sections.setdefault(current, [])
            continue
        sections[current].append(raw_line)
    return {k: "\n".join(v).strip() for k, v in sections.items() if "\n".join(v).strip()}


def _normalize_section_name(name: str) -> str:
    n = name.strip().lower()
    n = re.sub(r"[^a-z0-9\s/]", "", n)
    n = re.sub(r"\s+", " ", n)
    return n


def _merged_sections_from_collected_materials(items: list[CollectedMaterialIn]) -> dict[str, str]:
    """Merge Markdown sections from collected assistant outputs (no extra AI calls)."""
    core_map = {
        "quick summary": "Quick Summary",
        "key findings": "Key Findings",
        "recommended actions": "Recommended Actions",
        "references / citations": "References / Citations",
        "references citations": "References / Citations",
    }
    core_sections: dict[str, list[str]] = {v: [] for v in core_map.values()}
    extra_sections: dict[str, list[str]] = {}

    for item in items:
        text = (item.text or "").strip()
        if not text:
            continue
        parsed = _parse_markdown_sections(text)
        for heading, body in parsed.items():
            norm = _normalize_section_name(heading)
            target = core_map.get(norm)
            if target:
                core_sections[target].append(body)
            else:
                extra_sections.setdefault(heading.strip(), []).append(body)

    merged: dict[str, str] = {}
    for k in ("Quick Summary", "Key Findings", "Recommended Actions", "References / Citations"):
        joined = "\n\n".join([s for s in core_sections.get(k, []) if s.strip()]).strip()
        if joined:
            merged[k] = joined

    for heading, chunks in extra_sections.items():
        joined = "\n\n".join([s for s in chunks if s.strip()]).strip()
        if joined:
            merged[f"Additional: {heading}"] = joined

    return merged


def _reference_generated_sections() -> dict:
    """Structured report blocks constrained to project-scoped references."""
    return {
        "4.1": (
            "Space inventory from drawings: room labels, dimensions, symbols, and material callouts "
            "with stated measurement units (metric / imperial as declared on sheets)."
        ),
        "4.2": (
            "Gross vs usable areas: reconcile boundaries, mezzanines, and exclusions; surface assumptions "
            "where linework or scale is ambiguous."
        ),
        "4.3": (
            "Space Types Summary: gross, usable, counts, percentages, floor-wise breakdown, and "
            "efficiency ratio — all in the project’s stated measurement units."
        ),
        "4.4": {
            "chart_title": "Space types by area (illustrative)",
            "slices": [
                {"label": "Office", "area": 1240, "pct": 38},
                {"label": "Circulation", "area": 520, "pct": 16},
                {"label": "MEP / riser", "area": 310, "pct": 10},
                {"label": "Amenity", "area": 680, "pct": 21},
                {"label": "Other", "area": 480, "pct": 15},
            ],
        },
        "4.5": {
            "manual_reinspection_required": [
                {
                    "item": "North stair dimension extraction",
                    "doubt_score": 78,
                    "manual_review_reason": "Low OCR confidence on boundary text",
                },
                {
                    "item": "Ceiling height schedule vs reflected ceiling plan",
                    "doubt_score": 62,
                    "manual_review_reason": "Conflicting annotations between RCP and door schedule",
                },
            ]
        },
        "4.6": (
            "AI-generated suggestions for fixes: prioritized corrective actions tied to cited rules, "
            "drawing references, and re-measurement steps."
        ),
        "4.7": (
            "Pipeline note: regulations parsing → structured rules → automated compliance checks with "
            "citations strictly tied to Selected Ordinance Docs and drawing/project-file evidence."
        ),
        "compliance": {
            "status": "Structured rules vs project evidence",
            "citations": [
                {"section": "Selected Ordinance Doc: section reference", "drawing_page": "A-103"},
                {"section": "Selected Ordinance Doc: accessibility clause reference", "drawing_page": "G-201"},
            ],
        },
    }


def _report_payload(
    project_id: int,
    prompts: list[str] | None = None,
    chat_history: list[dict] | None = None,
    annotation_assets: list[dict] | None = None,
) -> dict:
    chat_rows = list(chat_history or [])
    generated: dict = {}

    if prompts:
        clean_prompts = [p.strip() for p in prompts if isinstance(p, str) and p.strip()]
        if clean_prompts:
            for i, prompt in enumerate(clean_prompts, start=1):
                reply = _assistant_reply_after_user_message(chat_rows, prompt)
                generated[f"Prompt {i}"] = (
                    f"User message:\n{prompt}\n\nAssistant reply:\n{reply}"
                )

    for key, val in _reference_generated_sections().items():
        generated[key] = val

    base = {
        "project_id": project_id,
        "generated_sections": generated,
        "citations": [
            {"source": "Selected Ordinance Docs + Drawings + Project Files", "confidence": 0.76, "doubt_score": 24}
        ],
        "annotated_preview_note": "Issue overlays should be shown in red on drawing previews.",
        "chart_metrics": {
            "title": "Review metrics (illustrative)",
            "bars": [
                {"label": "Confidence", "value": 76, "color": "#059669"},
                {"label": "Doubt / re-check", "value": 24, "color": "#d97706"},
                {"label": "Open items", "value": 12, "color": "#dc2626"},
            ],
        },
        "deficiency_annotation_samples": [
            {
                "drawing_sheet": "A-103",
                "annotation_type": "Deficiency",
                "description": "Stair width unclear at north core — verify against code min.",
                "severity": "High",
            },
            {
                "drawing_sheet": "M-201",
                "annotation_type": "Discrepancy",
                "description": "Diffuser count differs from schedule line 14.",
                "severity": "Medium",
            },
        ],
        "annotated_drawing_assets": annotation_assets or [],
        "chat_history": chat_rows,
    }
    if prompts:
        clean_prompts = [p.strip() for p in prompts if isinstance(p, str) and p.strip()]
        if clean_prompts:
            base["selected_user_prompts"] = clean_prompts
    return base


def _apply_report_metadata(payload: dict, project: Project, report_title: str, author: str | None) -> None:
    payload["report_title"] = report_title
    payload["project_name"] = project.name
    payload["author"] = (author or "").strip() or "—"
    payload["generated_at"] = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")


def _html_svg_bar_dashboard(chart: dict | None) -> str:
    if not chart or not chart.get("bars"):
        return ""
    title = escape(str(chart.get("title", "Metrics")))
    bars = chart["bars"]
    vals = [int(b.get("value", 0) or 0) for b in bars]
    max_v = max(vals) if vals else 1
    w = 640
    row_h = 40
    svg_h = 52 + len(bars) * row_h
    parts = [
        '<div class="dashboard">',
        f'<h2 class="section-title">{title}</h2>',
        f'<svg viewBox="0 0 {w} {svg_h}" xmlns="http://www.w3.org/2000/svg" role="img" aria-label="Bar chart" '
        f'style="font-family: {REPORT_FONT_STACK_XML}">',
    ]
    y = 44
    for b in bars:
        label = escape(str(b.get("label", "")))
        val = int(b.get("value", 0) or 0)
        color = escape(str(b.get("color", "#059669")))
        bw = int((val / max_v) * (w - 200)) if max_v else 0
        parts.append(f'<text x="8" y="{y + 22}" font-size="14" font-weight="600" fill="#1f2937">{label}</text>')
        parts.append(f'<rect x="168" y="{y - 6}" width="{max(bw, 6)}" height="28" fill="{color}" rx="5" opacity="0.9"/>')
        parts.append(f'<text x="{176 + bw}" y="{y + 12}" font-size="13" fill="#111827">{val}</text>')
        y += row_h
    parts.append("</svg></div>")
    return "\n".join(parts)


def _html_annotation_table(payload: dict) -> str:
    samples = payload.get("deficiency_annotation_samples") or []
    if not samples:
        return ""
    head = "<thead><tr><th>Drawing / sheet</th><th>Type</th><th>Description</th><th>Severity</th></tr></thead>"
    body_rows = []
    for row in samples:
        body_rows.append(
            "<tr>"
            f"<td>{escape(str(row.get('drawing_sheet', '')))}</td>"
            f"<td><span class='pill pill-{escape(str(row.get('annotation_type', '')).lower())}'>{escape(str(row.get('annotation_type', '')))}</span></td>"
            f"<td>{escape(str(row.get('description', '')))}</td>"
            f"<td>{escape(str(row.get('severity', '')))}</td>"
            "</tr>"
        )
    note = escape(str(payload.get("annotated_preview_note", "")))
    return f"""
    <section class="block">
      <h2 class="section-title">Deficiency &amp; discrepancy annotations</h2>
      <p class="lead">Mark issues on drawings with clear callouts (deficiency = red, discrepancy = amber). Example log:</p>
      <table class="data-table">{head}<tbody>{''.join(body_rows)}</tbody></table>
      <p class="muted small">{note}</p>
    </section>
    """


def _html_attached_annotations(payload: dict) -> str:
    assets = payload.get("annotated_drawing_assets") or []
    if not assets:
        return ""
    title = escape(str(payload.get("attached_annotations_section_title") or "Attached annotated drawings"))
    lead = escape(
        str(
            payload.get("attached_annotations_section_lead")
            or "User-attached drawing overlays included with this report export."
        )
    )
    cards: list[str] = []
    for i, row in enumerate(assets, start=1):
        filename = escape(str(row.get("filename", f"Drawing {i}")))
        marker_count = len(row.get("markers") or [])
        data_url = str(row.get("image_data_url", "")).strip()
        img_html = (
            f'<img src="{escape(data_url)}" alt="{filename}" class="annotated-preview" />'
            if data_url.startswith("data:image/")
            else '<div class="muted small">Preview unavailable</div>'
        )
        cards.append(
            f"""
            <div class="annotated-card">
              <div class="annotated-head"><strong>{filename}</strong><span class="muted small">{marker_count} marker(s)</span></div>
              {img_html}
            </div>
            """.strip()
        )
    return f"""
    <section class="block">
      <h2 class="section-title">{title}</h2>
      <p class="lead">{lead}</p>
      <div class="annotated-grid">{''.join(cards)}</div>
    </section>
    """


def _html_space_type_pie_section(gs: dict) -> str:
    raw = gs.get("4.4")
    if not isinstance(raw, dict):
        return ""
    title = escape(str(raw.get("chart_title", "Space types by area")))
    slices = raw.get("slices") or []
    if not isinstance(slices, list) or not slices:
        return ""
    colors = ["#10b981", "#0d9488", "#14b8a6", "#5eead4", "#94a3b8", "#64748b", "#334155"]
    cx, cy, r = 120, 120, 88

    def _pt(angle_deg: float) -> tuple[float, float]:
        rad = math.radians(angle_deg)
        return cx + r * math.cos(rad), cy + r * math.sin(rad)

    parts = [
        '<section class="block exec-card">',
        f'<h2 class="section-title">{title}</h2>',
        '<div class="pie-wrap"><svg viewBox="0 0 280 240" xmlns="http://www.w3.org/2000/svg" role="img" '
        f'aria-label="Pie chart" style="font-family: {REPORT_FONT_STACK_XML}">',
    ]
    start = -90.0
    total = sum(float(s.get("area", 0) or 0) for s in slices if isinstance(s, dict)) or 1.0
    for i, s in enumerate(slices):
        if not isinstance(s, dict):
            continue
        frac = float(s.get("area", 0) or 0) / total
        if frac <= 0:
            continue
        angle = frac * 360.0
        end = start + angle
        x1, y1 = _pt(start)
        x2, y2 = _pt(end)
        large = 1 if angle > 180 else 0
        color = colors[i % len(colors)]
        parts.append(
            f'<path d="M {cx} {cy} L {x1:.2f} {y1:.2f} A {r} {r} 0 {large} 1 {x2:.2f} {y2:.2f} Z" fill="{escape(color)}" stroke="#fff" stroke-width="1"/>'
        )
        start = end
    parts.append("</svg>")
    parts.append('<ul class="pie-legend">')
    for i, s in enumerate(slices):
        if not isinstance(s, dict):
            continue
        lbl = escape(str(s.get("label", "")))
        area = escape(str(s.get("area", "")))
        pct = escape(str(s.get("pct", "")))
        sw = colors[i % len(colors)]
        parts.append(
            f'<li><span class="swatch" style="background:{escape(sw)}"></span>'
            f"<strong>{lbl}</strong> — {area} ({pct}%)</li>"
        )
    parts.append("</ul></div></section>")
    return "\n".join(parts)


def _html_manual_reinspection_section(gs: dict) -> str:
    raw = gs.get("4.5")
    if not isinstance(raw, dict):
        return ""
    rows = raw.get("manual_reinspection_required") or []
    if not isinstance(rows, list) or not rows:
        return ""
    cards = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        item = escape(str(row.get("item", "")))
        score = escape(str(row.get("doubt_score", "")))
        reason = escape(str(row.get("manual_review_reason", "")))
        cards.append(
            f'<div class="reinspect-card"><div class="reinspect-top">'
            f'<strong>{item}</strong><span class="doubt-pill">doubt_score {score}</span></div>'
            f'<p class="reinspect-reason">{reason}</p></div>'
        )
    inner = "".join(cards)
    return f"""
    <section class="block reinspect-banner">
      <h2 class="reinspect-title">Manual Re-inspection Required</h2>
      <p class="lead reinspect-lead">Items below exceed confidence thresholds or have missing inputs. Re-verify on sheets before approval.</p>
      <div class="reinspect-grid">{inner}</div>
    </section>
    """


def _html_compliance_banner(gs: dict) -> str:
    comp = gs.get("compliance")
    if not isinstance(comp, dict):
        return ""
    status = escape(str(comp.get("status", "")))
    cites = comp.get("citations") or []
    if not isinstance(cites, list):
        cites = []
    cite_lines = []
    for c in cites[:12]:
        if not isinstance(c, dict):
            continue
        cite_lines.append(
            "<li>"
            f"<strong>{escape(str(c.get('section', '')))}</strong>"
            f" — sheet {escape(str(c.get('drawing_page', '')))}</li>"
        )
    ul = f"<ul class='cite-list'>{''.join(cite_lines)}</ul>" if cite_lines else ""
    return f"""
    <section class="block compliance-card">
      <h2 class="section-title">Ordinance comparison &amp; citations</h2>
      <p class="compliance-status"><strong>Status:</strong> {status}</p>
      {ul}
    </section>
    """


def _decode_png_data_url(data_url: str) -> bytes | None:
    prefix = "data:image/png;base64,"
    if not data_url.startswith(prefix):
        return None
    encoded = data_url[len(prefix) :]
    try:
        return base64.b64decode(encoded, validate=True)
    except (binascii.Error, ValueError):
        return None


def _write_report_files(report_id: str, payload: dict) -> dict:
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    pdf_path = REPORTS_DIR / f"{report_id}.pdf"
    docx_path = REPORTS_DIR / f"{report_id}.docx"
    html_path = REPORTS_DIR / f"{report_id}.html"
    # Layout hints for HTML export only (omit from stored JSON).
    attached_annotations_last = bool(payload.pop("attached_annotations_last", False))

    report_title = str(payload.get("report_title") or "AI Assistant Report")
    author = str(payload.get("author") or "—")
    project_name = str(payload.get("project_name") or f"Project #{payload.get('project_id', '')}")
    generated_at = str(payload.get("generated_at") or datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"))

    chat_rows = payload.get("chat_history") or []

    chat_blocks_html: list[str] = []
    for row in chat_rows:
        role = str(row.get("role", "assistant"))
        created_at = str(row.get("created_at", ""))
        text = str(row.get("text", ""))
        if role == "assistant":
            body_html = md.markdown(text, extensions=["tables", "fenced_code", "sane_lists"])
        else:
            body_html = f"<p>{escape(text).replace(chr(10), '<br/>')}</p>"
        label = "AI Assistant" if role == "assistant" else "Source"
        chat_blocks_html.append(
            f"""
            <article class="chat-item {escape(role)}">
              <div class="chat-role">{escape(label)} · {escape(created_at)}</div>
              <div class="reply">{body_html}</div>
            </article>
            """.strip()
        )

    gs = payload.get("generated_sections") or {}
    chart_html = _html_svg_bar_dashboard(payload.get("chart_metrics"))
    manual_html = _html_manual_reinspection_section(gs if isinstance(gs, dict) else {})
    comp_html = _html_compliance_banner(gs if isinstance(gs, dict) else {})
    pie_html = _html_space_type_pie_section(gs if isinstance(gs, dict) else {})
    featured_row = f'<div class="exec-row">{chart_html}{pie_html}</div>' if (chart_html or pie_html) else ""
    ann_html = _html_annotation_table(payload)
    attached_ann_html = _html_attached_annotations(payload)

    transcript_inner = (
        "".join(chat_blocks_html) if chat_blocks_html else '<p class="muted">No transcript included.</p>'
    )
    transcript_section = f"""    <section class="block transcript">
      <h2 class="section-title">Transcript &amp; narrative</h2>
      {transcript_inner}
    </section>"""
    if attached_annotations_last and attached_ann_html.strip():
        tail_sections = f"{ann_html}\n{transcript_section}\n{attached_ann_html}"
    else:
        tail_sections = f"{ann_html}\n{attached_ann_html}\n{transcript_section}"

    report_html = f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <title>{escape(REPORT_BRAND_HEADER)} — {escape(report_title)} — {escape(report_id)}</title>
  <style>
    :root {{ --report-font: {REPORT_FONT_STACK}; }}
    html, body {{ font-family: var(--report-font); margin: 0; color: #0f172a; background: #f8fafc; line-height: 1.55; }}
    .wrap, .wrap * {{ font-family: var(--report-font); }}
    svg text {{ font-family: var(--report-font); }}
    .wrap {{ max-width: 920px; margin: 0 auto; padding: 32px 24px 48px; background: #fff; min-height: 100vh;
      box-shadow: 0 1px 3px rgba(15,23,42,0.08); }}
    .report-header {{ border-bottom: 2px solid #10b981; padding-bottom: 20px; margin-bottom: 28px; }}
    .report-header h1 {{ margin: 0 0 4px; font-size: 1.75rem; font-weight: 700; letter-spacing: 0; }}
    .report-subhead {{ margin: 0 0 10px; font-size: 1.05rem; font-weight: 600; color: #334155; }}
    .meta {{ color: #64748b; font-size: 0.9rem; display: flex; flex-wrap: wrap; gap: 12px 24px; }}
    .meta strong {{ color: #334155; }}
    .section-title {{ font-size: 1.2rem; margin: 0 0 12px; color: #0f172a; border-left: 4px solid #10b981; padding-left: 10px; }}
    .lead {{ font-size: 1rem; margin: 0 0 12px; }}
    .muted {{ color: #64748b; }}
    .small {{ font-size: 0.85rem; }}
    .block {{ margin-bottom: 32px; }}
    .dashboard {{ background: #f1f5f9; border-radius: 12px; padding: 16px 20px; margin-bottom: 8px; }}
    .dashboard svg {{ width: 100%; height: auto; display: block; }}
    .data-table {{ width: 100%; border-collapse: collapse; font-size: 0.9rem; margin: 12px 0; }}
    .data-table th, .data-table td {{ border: 1px solid #e2e8f0; padding: 10px 12px; text-align: left; vertical-align: top; }}
    .data-table thead th {{ background: #d1fae5; font-weight: 600; }}
    .pill {{ display: inline-block; padding: 2px 10px; border-radius: 999px; font-size: 0.8rem; font-weight: 600; }}
    .pill-deficiency {{ background: #fee2e2; color: #991b1b; }}
    .pill-discrepancy {{ background: #ffedd5; color: #9a3412; }}
    .transcript h2 {{ margin-top: 0; }}
    .chat-item {{ border: 1px solid #e2e8f0; border-radius: 12px; padding: 14px 16px; margin-bottom: 12px; background: #fff; }}
    .chat-item.user {{ border-color: #10b981; background: #f0fdf4; }}
    .chat-role {{ font-size: 0.75rem; color: #64748b; margin-bottom: 8px; text-transform: uppercase; letter-spacing: 0; }}
    .reply {{ font-size: 0.95rem; }}
    .reply h1 {{ font-size: 1.35rem; }}
    .reply h2 {{ font-size: 1.15rem; }}
    .reply h3 {{ font-size: 1.05rem; }}
    .reply p {{ margin: 6px 0 10px; }}
    .reply ul, .reply ol {{ margin: 6px 0 10px 22px; }}
    .reply table {{ width: 100%; border-collapse: collapse; margin: 10px 0 14px; font-size: 0.88rem; }}
    .reply th, .reply td {{ border: 1px solid #e2e8f0; padding: 8px 10px; }}
    .reply th {{ background: #f1f5f9; }}
    .footnote {{ margin-top: 40px; padding-top: 16px; border-top: 1px solid #e2e8f0; font-size: 0.8rem; color: #94a3b8; }}
    .annotated-grid {{ display: grid; gap: 14px; grid-template-columns: repeat(auto-fit, minmax(220px, 1fr)); }}
    .annotated-card {{ border: 1px solid #e2e8f0; border-radius: 12px; padding: 10px; background: #fff; }}
    .annotated-head {{ display: flex; justify-content: space-between; gap: 8px; margin-bottom: 8px; align-items: baseline; }}
    .annotated-preview {{ width: 100%; height: auto; display: block; border-radius: 8px; border: 1px solid #e2e8f0; }}
    .exec-row {{ display: flex; flex-wrap: wrap; gap: 20px; align-items: flex-start; margin-bottom: 8px; }}
    .exec-row .dashboard, .exec-row .exec-card {{ flex: 1 1 280px; min-width: 0; }}
    .pie-wrap {{ display: flex; flex-wrap: wrap; gap: 16px; align-items: center; }}
    .pie-legend {{ list-style: none; margin: 0; padding: 0; font-size: 0.88rem; }}
    .pie-legend li {{ margin: 6px 0; display: flex; align-items: center; gap: 8px; }}
    .swatch {{ width: 12px; height: 12px; border-radius: 3px; flex-shrink: 0; }}
    .reinspect-banner {{ background: linear-gradient(135deg, #fef2f2 0%, #fff7ed 100%); border: 1px solid #fecaca; border-radius: 14px; padding: 18px 20px; }}
    .reinspect-title {{ margin: 0 0 8px; font-size: 1.35rem; color: #991b1b; letter-spacing: 0; }}
    .reinspect-lead {{ margin: 0 0 14px; }}
    .reinspect-grid {{ display: grid; gap: 12px; grid-template-columns: repeat(auto-fit, minmax(240px, 1fr)); }}
    .reinspect-card {{ background: #fff; border: 1px solid #fecaca; border-radius: 10px; padding: 12px 14px; }}
    .reinspect-top {{ display: flex; justify-content: space-between; gap: 10px; align-items: flex-start; flex-wrap: wrap; }}
    .doubt-pill {{ font-size: 0.72rem; font-weight: 700; background: #fee2e2; color: #b91c1c; padding: 4px 10px; border-radius: 999px; white-space: nowrap; }}
    .reinspect-reason {{ margin: 8px 0 0; font-size: 0.88rem; color: #334155; }}
    .compliance-card {{ background: #f8fafc; border-radius: 12px; padding: 16px 18px; border: 1px solid #e2e8f0; }}
    .compliance-status {{ margin: 0 0 10px; }}
    .cite-list {{ margin: 0; padding-left: 1.1rem; }}
  </style>
</head>
<body>
  <div class="wrap">
    <header class="report-header">
      <h1>{escape(REPORT_BRAND_HEADER)}</h1>
      <p class="report-subhead">{escape(report_title)}</p>
      <div class="meta">
        <span><strong>Report ID</strong> {escape(report_id)}</span>
        <span><strong>Project</strong> {escape(project_name)}</span>
        <span><strong>Author</strong> {escape(author)}</span>
        <span><strong>Generated</strong> {escape(generated_at)}</span>
      </div>
    </header>
    {manual_html}
    {comp_html}
    {featured_row}
{tail_sections}
    <p class="footnote">CIA Construction Insight Agent — exported document. References are restricted to project documents only (Drawings, Project Files, Selected Ordinance Docs).</p>
  </div>
</body>
</html>
"""
    html_path.write_text(report_html, encoding="utf-8")

    annotated_assets = payload.get("annotated_drawing_assets") or []
    annotated_downloads: dict[str, str] = {}
    for i, asset in enumerate(annotated_assets, start=1):
        data_url = str(asset.get("image_data_url", "")).strip()
        image_bytes = _decode_png_data_url(data_url)
        if not image_bytes:
            continue
        stem = Path(str(asset.get("filename", f"drawing-{i}"))).stem
        clean_stem = "".join(ch if ch.isalnum() or ch in ("-", "_") else "-" for ch in stem).strip("-") or f"drawing-{i}"
        ann_name = f"{report_id}-annotated-{i}-{clean_stem}.png"
        ann_path = REPORTS_DIR / ann_name
        ann_path.write_bytes(image_bytes)
        annotated_downloads[f"annotated_png_{i}"] = f"/api/ai/reports/download/{ann_name}"

    fb_html = (
        f"<html><head><style>body{{font-family:{REPORT_FONT_STACK};}}</style></head>"
        f"<body><h1>{escape(report_title)}</h1>"
        "<p>PDF conversion failed; use the HTML export for the full report.</p></body></html>"
    )
    try:
        pdf_path.write_bytes(_html_to_pdf_bytes_for_report(report_html))
    except Exception:
        pdf_path.write_bytes(_html_to_pdf_bytes_xhtml2pdf(fb_html))

    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Pt
    from htmldocx import HtmlToDocx

    try:
        document = Document()
        _apply_report_docx_base_font(document)
        HtmlToDocx().add_html_to_document(report_html, document)
        document.save(str(docx_path))
    except Exception:
        docx_doc = Document()
        _apply_report_docx_base_font(docx_doc)
        t_para = docx_doc.add_paragraph()
        t_run = t_para.add_run(REPORT_BRAND_HEADER)
        t_run.bold = True
        t_run.font.size = Pt(22)
        t_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        sub = docx_doc.add_paragraph()
        sub_run = sub.add_run(report_title)
        sub_run.bold = True
        sub_run.font.size = Pt(14)
        docx_doc.add_paragraph(
            "Word export could not convert the full HTML layout automatically. "
            "Please use the HTML or PDF download for the complete document."
        )
        docx_doc.save(docx_path)

    return {
        "pdf": f"/api/ai/reports/download/{pdf_path.name}",
        "word": f"/api/ai/reports/download/{docx_path.name}",
        "html": f"/api/ai/reports/download/{html_path.name}",
        **annotated_downloads,
    }


def _build_consolidated_collection_report(
    project: Project,
    body: ConsolidateCollectionReportRequest,
) -> tuple[str, dict, dict[str, str]]:
    """Assemble Custom Report from collected materials only (no per-item OpenRouter calls)."""
    report_id = f"consolidated-collection-{body.project_id}-{uuid4().hex[:8]}"
    now_iso = datetime.now(timezone.utc).isoformat()
    transcript: list[dict] = []
    for item in body.items:
        t = (item.text or "").strip()
        if not t:
            continue
        transcript.append(
            {
                "role": "user",
                "text": f"{item.label} ({item.source})",
                "created_at": item.created_at or now_iso,
            }
        )
        transcript.append({"role": "assistant", "text": t, "created_at": item.created_at or now_iso})
    if not transcript:
        for item in body.items:
            header = f"### {item.label}\n_Source: {item.source}_ · _ID: {item.id or '—'}_\n\n"
            transcript.append(
                {
                    "role": "assistant",
                    "text": header + (item.text or ""),
                    "created_at": item.created_at or now_iso,
                }
            )

    ann_payload = (
        [a.model_dump() for a in body.annotation_assets] if body.include_annotated_drawing else []
    )
    payload = _report_payload(
        body.project_id,
        chat_history=transcript,
        annotation_assets=ann_payload,
    )
    merged_sections = _merged_sections_from_collected_materials(body.items)
    collected_gs: dict[str, str] = {}
    for idx, item in enumerate(body.items, start=1):
        label = (item.label or f"Output {idx}").strip()
        if len(label) > 76:
            label = f"{label[:73]}..."
        collected_gs[f"Collected {idx}: {label}"] = (
            f"Source: {item.source}\nID: {item.id or '—'}\nCreated: {item.created_at or '—'}\n\n{item.text}"
        )
    payload["generated_sections"] = {**merged_sections, **collected_gs, **payload["generated_sections"]}
    payload["collected_materials"] = [m.model_dump() for m in body.items]
    if body.include_annotated_drawing and ann_payload:
        payload["attached_annotations_last"] = True
        payload["attached_annotations_section_title"] = "Annotated drawing — Custom Report"
        payload["attached_annotations_section_lead"] = (
            "Annotated image(s) from Drawings or Project files (final section of this Custom Report)."
        )
    _apply_report_metadata(payload, project, "Consolidated Report (Collected Outputs)", body.author)
    files = _write_report_files(report_id, payload)
    return report_id, payload, files


def _run_custom_report_job_task(job_id: int, body_dict: dict[str, Any], owner_id: int) -> None:
    db = SessionLocal()
    job: ReportJob | None = None
    try:
        body = ConsolidateCollectionReportRequest(**body_dict)
        job = db.query(ReportJob).filter(ReportJob.id == job_id, ReportJob.owner_id == owner_id).first()
        if not job:
            return
        job.status = "running"
        db.commit()
        project = db.query(Project).filter(Project.id == body.project_id, Project.owner_id == owner_id).first()
        if not project:
            job.status = "failed"
            job.error_message = "Project not found"
            job.completed_at = datetime.now(timezone.utc)
            db.commit()
            return
        report_id, _payload, files = _build_consolidated_collection_report(project, body)
        job.status = "completed"
        job.report_id = report_id
        job.downloads_json = json.dumps(files)
        job.completed_at = datetime.now(timezone.utc)
        db.commit()
    except Exception as e:
        db.rollback()
        job = db.query(ReportJob).filter(ReportJob.id == job_id).first()
        if job:
            job.status = "failed"
            job.error_message = (str(e) or "Report job failed")[:4000]
            job.completed_at = datetime.now(timezone.utc)
            db.commit()
    finally:
        db.close()


def _build_standard_report(
    project: Project,
    body: ConsolidatedReportRequest,
    db: Session,
) -> tuple[str, dict, dict[str, str]]:
    report_id = f"standard-{body.project_id}-{uuid4().hex[:8]}"
    selected_prompts = [p.strip() for p in body.user_prompts if isinstance(p, str) and p.strip()]
    if not selected_prompts:
        for row in body.chat_history:
            if str(row.get("role", "")).lower() != "user":
                continue
            text = str(row.get("text", "")).strip()
            if text:
                selected_prompts.append(text)
    generated_transcript = _run_selected_prompts(selected_prompts, project, db)
    payload = _report_payload(
        body.project_id,
        prompts=selected_prompts,
        chat_history=generated_transcript or body.chat_history,
        annotation_assets=body.annotation_assets,
    )
    _apply_report_metadata(payload, project, "Standard AI Assistant Report", body.author)
    files = _write_report_files(report_id, payload)
    return report_id, payload, files


def _run_standard_report_job_task(job_id: int, body_dict: dict[str, Any], owner_id: int) -> None:
    db = SessionLocal()
    job: ReportJob | None = None
    try:
        body = ConsolidatedReportRequest(**body_dict)
        job = db.query(ReportJob).filter(ReportJob.id == job_id, ReportJob.owner_id == owner_id).first()
        if not job:
            return
        job.status = "running"
        db.commit()
        project = db.query(Project).filter(Project.id == body.project_id, Project.owner_id == owner_id).first()
        if not project:
            job.status = "failed"
            job.error_message = "Project not found"
            job.completed_at = datetime.now(timezone.utc)
            db.commit()
            return
        report_id, _payload, files = _build_standard_report(project, body, db)
        job.status = "completed"
        job.report_id = report_id
        job.downloads_json = json.dumps(files)
        job.completed_at = datetime.now(timezone.utc)
        db.commit()
    except Exception as e:
        db.rollback()
        job = db.query(ReportJob).filter(ReportJob.id == job_id).first()
        if job:
            job.status = "failed"
            job.error_message = (str(e) or "Standard report job failed")[:4000]
            job.completed_at = datetime.now(timezone.utc)
            db.commit()
    finally:
        db.close()


@router.post("/reports/standard/jobs", response_model=StandardReportJobAccepted)
def enqueue_standard_report(
    body: ConsolidatedReportRequest,
    background_tasks: BackgroundTasks,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    project = db.query(Project).filter(Project.id == body.project_id, Project.owner_id == user.id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    job = ReportJob(
        owner_id=user.id,
        project_id=body.project_id,
        kind="standard_report",
        status="pending",
    )
    db.add(job)
    db.commit()
    db.refresh(job)
    background_tasks.add_task(_run_standard_report_job_task, job.id, body.model_dump(), user.id)
    return StandardReportJobAccepted(job_id=job.id, status="pending")


@router.get("/reports/standard/jobs/{job_id}", response_model=StandardReportJobStatusOut)
def get_standard_report_job_status(
    job_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    job = (
        db.query(ReportJob)
        .filter(ReportJob.id == job_id, ReportJob.owner_id == user.id, ReportJob.kind == "standard_report")
        .first()
    )
    if not job:
        raise HTTPException(status_code=404, detail="Report job not found")
    downloads = json.loads(job.downloads_json) if job.downloads_json else None
    return StandardReportJobStatusOut(
        job_id=job.id,
        status=job.status,
        report_type="standard" if job.status == "completed" else None,
        report_id=job.report_id,
        downloads=downloads if isinstance(downloads, dict) else None,
        error_message=job.error_message,
    )


@router.post("/reports/standard")
def generate_standard_report(
    body: ConsolidatedReportRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    project = db.query(Project).filter(Project.id == body.project_id, Project.owner_id == user.id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    report_id, payload, files = _build_standard_report(project, body, db)
    return {"report_type": "standard", "report_id": report_id, "payload": payload, "downloads": files}


@router.post("/reports/consolidated")
def generate_consolidated_report(
    body: ConsolidatedReportRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    project = db.query(Project).filter(Project.id == body.project_id, Project.owner_id == user.id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    report_id = f"consolidated-{body.project_id}-{uuid4().hex[:8]}"
    payload = _report_payload(
        body.project_id,
        body.user_prompts,
        body.chat_history,
        body.annotation_assets,
    )
    _apply_report_metadata(payload, project, "Consolidated AI Assistant Report", body.author)
    files = _write_report_files(report_id, payload)
    return {"report_type": "consolidated", "report_id": report_id, "payload": payload, "downloads": files}


@router.post("/reports/consolidate-collection/jobs", response_model=CustomReportJobAccepted)
def enqueue_consolidate_collection_report(
    body: ConsolidateCollectionReportRequest,
    background_tasks: BackgroundTasks,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    project = db.query(Project).filter(Project.id == body.project_id, Project.owner_id == user.id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    job = ReportJob(
        owner_id=user.id,
        project_id=body.project_id,
        kind="custom_report",
        status="pending",
    )
    db.add(job)
    db.commit()
    db.refresh(job)
    background_tasks.add_task(_run_custom_report_job_task, job.id, body.model_dump(), user.id)
    return CustomReportJobAccepted(job_id=job.id, status="pending")


@router.get("/reports/consolidate-collection/jobs/{job_id}", response_model=CustomReportJobStatusOut)
def get_consolidate_collection_job_status(
    job_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    job = db.query(ReportJob).filter(ReportJob.id == job_id, ReportJob.owner_id == user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Report job not found")
    downloads = json.loads(job.downloads_json) if job.downloads_json else None
    return CustomReportJobStatusOut(
        job_id=job.id,
        status=job.status,
        report_type="consolidated_collection" if job.status == "completed" else None,
        report_id=job.report_id,
        downloads=downloads if isinstance(downloads, dict) else None,
        error_message=job.error_message,
    )


@router.post("/reports/consolidate-collection")
def generate_consolidate_collection_report(
    body: ConsolidateCollectionReportRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Synchronous Custom Report (same assembly as background job). Prefer /jobs for long runs."""
    project = db.query(Project).filter(Project.id == body.project_id, Project.owner_id == user.id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    report_id, payload, files = _build_consolidated_collection_report(project, body)
    return {"report_type": "consolidated_collection", "report_id": report_id, "payload": payload, "downloads": files}


@router.get("/reports/download/{filename}")
def download_report_file(filename: str, _: User = Depends(get_current_user)):
    path = REPORTS_DIR / filename
    if not path.exists() or not path.is_file():
        raise HTTPException(status_code=404, detail="Report file not found")
    return FileResponse(path, media_type="application/octet-stream", filename=path.name)
